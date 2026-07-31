"""Text extraction from uploaded documents.

Extraction failures raise ``TextExtractionError``. They used to be *returned
as the document's text*, so a document that could not be parsed was recorded
as completed with the body "Error extracting text: ...". That string was then
chunked, embedded and indexed, and would come back as a search result and be
cited in an answer. A failure has to look like a failure.
"""

from __future__ import annotations

import logging
import os
import re
from datetime import datetime
from typing import Any, Dict, Tuple

import aiofiles

from backend.common.ocr import (
    MIN_CHARS_FOR_TEXT_LAYER,
    ocr_availability,
    ocr_image_file,
    ocr_pdf_pages,
)

try:
    import pdfplumber
except ImportError:  # pragma: no cover - optional dependency
    pdfplumber = None

try:
    import PyPDF2
except ImportError:  # pragma: no cover - optional dependency
    PyPDF2 = None

try:
    from docx import Document as DocxDocument
except ImportError:  # pragma: no cover - optional dependency
    DocxDocument = None

logger = logging.getLogger(__name__)

#: Extensions handled by a text extractor.
TEXT_EXTENSIONS = {".txt", ".md", ".csv"}
#: Image formats that need OCR, which is not wired into the pipeline yet.
IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".tiff", ".tif"}


class TextExtractionError(Exception):
    """Raised when a document's text cannot be extracted.

    The message is written for the person who uploaded the file and is
    surfaced verbatim on the document record.
    """


async def extract_text_from_file(file_path: str) -> Tuple[str, Dict[str, Any]]:
    """Extract text and metadata from a document.

    Args:
        file_path: Absolute path to the stored upload.

    Returns:
        ``(text, metadata)``. ``text`` is always real document content.

    Raises:
        TextExtractionError: if the format is unsupported, the file is
            corrupt, or it contains no extractable text.
    """
    file_extension = os.path.splitext(file_path)[1].lower()

    metadata: Dict[str, Any] = {
        "filename": os.path.basename(file_path),
        "file_type": file_extension,
        "file_size": os.path.getsize(file_path),
        "extracted_at": datetime.utcnow().isoformat() + "Z",
    }

    if file_extension == ".pdf":
        text, extra = await _extract_from_pdf(file_path)
    elif file_extension in (".doc", ".docx"):
        text, extra = await _extract_from_docx(file_path)
    elif file_extension in TEXT_EXTENSIONS:
        text, extra = await _extract_from_text(file_path)
    elif file_extension in IMAGE_EXTENSIONS:
        text, extra = await _extract_from_image(file_path)
    else:
        raise TextExtractionError(
            f"'{file_extension}' is not a supported document format."
        )

    metadata.update(extra)

    if not text.strip():
        raise TextExtractionError(
            "No text could be read from this document. It may be a scan "
            "without a text layer, password-protected, or empty."
        )

    return text, metadata


async def _extract_from_text(file_path: str) -> Tuple[str, Dict[str, Any]]:
    """Read a plain-text file, tolerating imperfect encodings."""
    try:
        async with aiofiles.open(file_path, encoding="utf-8") as handle:
            text = await handle.read()
    except UnicodeDecodeError:
        # Latin-1 maps every byte, so this always succeeds; better a few odd
        # characters than rejecting a document over one bad byte.
        logger.info("Falling back to latin-1 for %s", file_path)
        async with aiofiles.open(file_path, encoding="latin-1") as handle:
            text = await handle.read()

    return text, {"pages": 1, "word_count": len(text.split())}


async def _extract_from_pdf(file_path: str) -> Tuple[str, Dict[str, Any]]:
    """Extract text from a PDF, page by page.

    Page markers are emitted so retrieval can map a chunk back to a page
    number and cite it.
    """
    if pdfplumber is None and PyPDF2 is None:
        raise TextExtractionError(
            "No PDF library is installed on the server. "
            "Install pdfplumber to enable PDF support."
        )

    metadata: Dict[str, Any] = {"pages": 0, "word_count": 0}
    # Page number -> text, so pages without a text layer can be identified and
    # re-read with OCR before the document is assembled.
    pages: Dict[int, str] = {}
    failed_pages = []

    try:
        if pdfplumber is not None:
            with pdfplumber.open(file_path) as pdf:
                metadata["pages"] = len(pdf.pages)
                for page_number, page in enumerate(pdf.pages, start=1):
                    try:
                        pages[page_number] = page.extract_text() or ""
                    except Exception as exc:
                        # One unreadable page should not lose the other 99.
                        logger.warning(
                            "Page %s of %s could not be read: %s",
                            page_number,
                            file_path,
                            exc,
                        )
                        failed_pages.append(page_number)
        else:
            with open(file_path, "rb") as handle:
                reader = PyPDF2.PdfReader(handle)
                metadata["pages"] = len(reader.pages)
                for page_number, page in enumerate(reader.pages, start=1):
                    try:
                        pages[page_number] = page.extract_text() or ""
                    except Exception as exc:
                        logger.warning(
                            "Page %s of %s could not be read: %s",
                            page_number,
                            file_path,
                            exc,
                        )
                        failed_pages.append(page_number)

    except TextExtractionError:
        raise
    except Exception as exc:
        raise TextExtractionError(
            f"This PDF could not be opened: {exc}. It may be corrupt or "
            "password-protected."
        ) from exc

    # A page whose text layer is empty or near-empty is a scan. Read it with
    # OCR rather than letting it contribute nothing, which previously made a
    # scanned document either fail outright or index as blank.
    scanned = [
        number
        for number, body in sorted(pages.items())
        if len(body.strip()) < MIN_CHARS_FOR_TEXT_LAYER
    ]
    if scanned:
        available, reason = ocr_availability()
        if available:
            recovered = await ocr_pdf_pages(file_path, scanned)
            for page_number, page_text in recovered.items():
                pages[page_number] = page_text
            if recovered:
                # Surfaced on the document so a reader knows which pages were
                # machine-read rather than extracted, and can weigh them.
                metadata["ocr_pages"] = sorted(recovered)
                metadata["ocr_engine"] = reason
        else:
            metadata["ocr_unavailable_reason"] = reason
            metadata["pages_without_text_layer"] = scanned

    text = _normalise_whitespace(
        "\n".join(
            f"--- Page {number} ---\n{body}"
            for number, body in sorted(pages.items())
            if body.strip()
        )
    )
    metadata["word_count"] = len(text.split())
    if failed_pages:
        # Recorded on the document so a partial result is visible rather than
        # passing as complete.
        metadata["failed_pages"] = failed_pages

    return text, metadata


async def _extract_from_image(file_path: str) -> Tuple[str, Dict[str, Any]]:
    """Read text off an image with OCR.

    When OCR is unavailable the upload is still refused, with the reason it is
    unavailable rather than a blanket "not supported". Indexing a placeholder
    would make the document appear searchable while holding nothing.
    """
    available, reason = ocr_availability()
    if not available:
        raise TextExtractionError(
            f"This is an image, so it needs OCR to be read, and OCR is not "
            f"available: {reason}"
        )

    text = await ocr_image_file(file_path)
    if not text or not text.strip():
        raise TextExtractionError(
            "OCR ran on this image but found no readable text. It may be a "
            "photo without writing, or too low-resolution to read."
        )

    # Page 1 marker keeps citations uniform with PDFs, so an image cites as
    # "page 1" rather than as a special case everywhere downstream.
    return f"--- Page 1 ---\n{text}", {
        "pages": 1,
        "word_count": len(text.split()),
        "ocr_pages": [1],
        "ocr_engine": reason,
    }


async def _extract_from_docx(file_path: str) -> Tuple[str, Dict[str, Any]]:
    """Extract paragraphs and table cells from a Word document."""
    if DocxDocument is None:
        raise TextExtractionError(
            "python-docx is not installed on the server, so Word documents "
            "cannot be read."
        )

    try:
        document = DocxDocument(file_path)
    except Exception as exc:
        raise TextExtractionError(
            f"This Word document could not be opened: {exc}. Note that the "
            "legacy .doc format is not supported — save it as .docx."
        ) from exc

    parts = [p.text for p in document.paragraphs if p.text.strip()]

    # Tables carry much of the content in contracts and invoices; taking only
    # paragraphs silently dropped it.
    table_count = 0
    for table in document.tables:
        table_count += 1
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    text = _normalise_whitespace("\n".join(parts))
    return text, {
        "pages": 1,
        "paragraphs": len(document.paragraphs),
        "tables": table_count,
        "word_count": len(text.split()),
    }


#: pdfplumber emits (cid:NNN) for glyphs it cannot map to a character —
#: bullets and dingbats, mostly. Left in place they end up in citation
#: previews and in the context handed to the model.
_CID_ARTEFACT = re.compile(r"\(cid:\d+\)")


def _normalise_whitespace(text: str) -> str:
    """Clean up extraction artefacts and collapse redundant whitespace."""
    text = _CID_ARTEFACT.sub("•", text)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()
