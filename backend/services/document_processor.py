"""
Enhanced document processing service with local LLM and document parsing
"""

import os
import uuid
import aiofiles
from typing import Dict, Any, Optional, Callable
from datetime import datetime
import logging
import asyncio

# Document processing imports
import PyPDF2
from docx import Document as DocxDocument
from PIL import Image
import io

from backend.core.config import get_settings
from backend.services.local_llm_service import LocalLLMService
from backend.services.retrieval_service import RetrievalService

logger = logging.getLogger(__name__)


class DocumentProcessorService:
    """Enhanced document processing with local LLM"""

    def __init__(self):
        self.settings = get_settings()
        self.llm_service = LocalLLMService()
        self.retrieval_service = RetrievalService()  # For RAG indexing
        self.processing_status: Dict[str, Dict[str, Any]] = {}

    async def initialize(self):
        """Initialize the document processor and LLM service"""
        try:
            await self.llm_service.initialize()
            await self.retrieval_service.initialize()  # Init vector store
        except Exception as e:
            logger.warning(f"⚠️ Document processor initialized without LLM: {e}")

        if self.llm_service.is_initialized:
            logger.info("✅ Document processor with local LLM initialized")
        else:
            logger.warning(
                "⚠️ Document processor initialized in BASIC mode (No AI features)"
            )

    async def save_uploaded_file(self, file, document_id: str, user_id: str) -> str:
        """Save uploaded file securely"""
        user_upload_dir = os.path.join(self.settings.upload_directory, user_id)
        os.makedirs(user_upload_dir, exist_ok=True)

        file_extension = os.path.splitext(file.filename)[1]
        secure_filename = f"{document_id}_{uuid.uuid4().hex[:8]}{file_extension}"
        file_path = os.path.join(user_upload_dir, secure_filename)

        async with aiofiles.open(file_path, "wb") as buffer:
            content = await file.read()
            await buffer.write(content)

        logger.info(f"File saved: {file_path}")
        return file_path

    def initialize_status(self, document_id: str, filename: str):
        """Initialize document status to queued"""
        self.processing_status[document_id] = {
            "document_id": document_id,
            "status": "queued",
            "progress": 0,
            "message": "Queued for processing",
            "filename": filename,
            "created_at": datetime.utcnow(),
        }

    async def get_document_status(self, document_id: str) -> Optional[Dict[str, Any]]:
        """Get document processing status"""
        return self.processing_status.get(document_id)

    async def process_document_with_ai(
        self,
        document_id: str,
        file_path: str,
        options: Dict[str, Any],
        filename: str = "unknown",
        progress_callback: Optional[Callable] = None,
    ):
        """Process document with local LLM analysis"""
        try:
            # Update to processing
            if document_id in self.processing_status:
                self.processing_status[document_id].update(
                    {
                        "status": "processing",
                        "progress": 10,
                        "message": "Extracting text",
                    }
                )
            else:
                # Fallback if not initialized (should not happen with new flow)
                self.processing_status[document_id] = {
                    "document_id": document_id,
                    "status": "processing",
                    "progress": 10,
                    "message": "Extracting text",
                    "filename": filename,
                    "created_at": datetime.utcnow(),
                }

            if progress_callback:
                await progress_callback(self.processing_status[document_id])

            # Stage 1: Extract text from document
            extracted_text, doc_metadata = await self._extract_text_from_file(file_path)

            await self._update_progress(
                document_id, 40, "Text extracted, analyzing with AI", progress_callback
            )

            # Stage 2: Enhance with local LLM
            if options.get("enhance_with_ai", True) and extracted_text.strip():
                enhanced_result = await self.llm_service.enhance_document_data(
                    extracted_text, doc_metadata
                )
            else:
                enhanced_result = {
                    "enhanced_text": extracted_text,
                    "summary": "Document text extracted successfully",
                    "key_points": [],
                    "entities": [],
                    "confidence": 0.8,
                }

            await self._update_progress(
                document_id,
                80,
                "AI analysis complete, indexing for search",
                progress_callback,
            )

            # Stage 3: Index document content for RAG retrieval
            try:
                indexing_result = await self.retrieval_service.index_document(
                    document_id=document_id,
                    content=extracted_text,
                    metadata={
                        "filename": filename,
                        **doc_metadata,
                        "summary": enhanced_result.get("summary", ""),
                    },
                )
                logger.info(
                    f"Indexed document {document_id}: {indexing_result.chunks_indexed} chunks"
                )
            except Exception as e:
                logger.warning(f"Failed to index document {document_id}: {e}")

            await self._update_progress(
                document_id, 95, "Finalizing", progress_callback
            )

            # Stage 3: Finalize results
            final_result = {
                "document_id": document_id,
                "status": "completed",
                "progress": 100,
                "message": "Complete",
                "filename": filename,
                "result": {
                    "original_text": extracted_text,
                    "metadata": doc_metadata,
                    "ai_analysis": enhanced_result,
                    "processing_options": options,
                },
                "created_at": datetime.utcnow(),
                "completed_at": datetime.utcnow(),
            }

            self.processing_status[document_id] = final_result

            if progress_callback:
                await progress_callback(final_result)

        except Exception as e:
            logger.error(f"Processing failed for {document_id}: {e}")
            self.processing_status[document_id] = {
                "document_id": document_id,
                "status": "failed",
                "error": str(e),
                "created_at": datetime.utcnow(),
            }

            if progress_callback:
                await progress_callback(self.processing_status[document_id])

        finally:
            # Cleanup uploaded file
            if os.path.exists(file_path):
                try:
                    os.remove(file_path)
                    logger.info(f"Cleaned up file: {file_path}")
                except Exception as e:
                    logger.warning(f"Failed to cleanup file {file_path}: {e}")

    async def _extract_text_from_file(
        self, file_path: str
    ) -> tuple[str, Dict[str, Any]]:
        """Extract text from various file formats"""
        file_extension = os.path.splitext(file_path)[1].lower()
        filename = os.path.basename(file_path)

        metadata = {
            "filename": filename,
            "file_type": file_extension,
            "file_size": os.path.getsize(file_path),
            "processing_time": datetime.utcnow().isoformat(),
        }

        try:
            if file_extension == ".pdf":
                text, pdf_metadata = await self._extract_from_pdf(file_path)
                metadata.update(pdf_metadata)
                return text, metadata

            elif file_extension in [".doc", ".docx"]:
                text, doc_metadata = await self._extract_from_docx(file_path)
                metadata.update(doc_metadata)
                return text, metadata

            elif file_extension == ".txt":
                async with aiofiles.open(file_path, "r", encoding="utf-8") as f:
                    text = await f.read()
                metadata.update({"pages": 1, "word_count": len(text.split())})
                return text, metadata

            elif file_extension in [".png", ".jpg", ".jpeg"]:
                # For images, you could add OCR here using pytesseract
                metadata.update({"type": "image", "ocr_available": False})
                return (
                    "Image file uploaded - OCR not implemented in this demo",
                    metadata,
                )

            else:
                return f"Unsupported file type: {file_extension}", metadata

        except Exception as e:
            logger.error(f"Text extraction failed for {file_path}: {e}")
            return f"Error extracting text: {str(e)}", metadata

    async def _extract_from_pdf(self, file_path: str) -> tuple[str, Dict[str, Any]]:
        """Extract text from PDF file"""
        text = ""
        metadata = {"pages": 0, "word_count": 0}

        try:
            with open(file_path, "rb") as file:
                pdf_reader = PyPDF2.PdfReader(file)
                metadata["pages"] = len(pdf_reader.pages)

                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        text += f"\n--- Page {page_num + 1} ---\n{page_text}\n"
                    except Exception as e:
                        logger.warning(
                            f"Failed to extract text from page {page_num + 1}: {e}"
                        )
                        text += (
                            f"\n--- Page {page_num + 1} ---\n[Text extraction failed]\n"
                        )

                metadata["word_count"] = len(text.split())
                return text.strip(), metadata

        except Exception as e:
            logger.error(f"PDF extraction failed: {e}")
            return f"PDF extraction error: {str(e)}", metadata

    async def _extract_from_docx(self, file_path: str) -> tuple[str, Dict[str, Any]]:
        """Extract text from DOCX file"""
        try:
            doc = DocxDocument(file_path)
            text = ""

            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"

            metadata = {
                "pages": 1,  # DOCX doesn't have clear page boundaries
                "paragraphs": len(doc.paragraphs),
                "word_count": len(text.split()),
            }

            return text.strip(), metadata

        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
            return f"DOCX extraction error: {str(e)}", {"pages": 0, "word_count": 0}

    async def _update_progress(
        self,
        document_id: str,
        progress: int,
        stage: str,
        callback: Optional[Callable] = None,
    ):
        """Update processing progress"""
        if document_id in self.processing_status:
            self.processing_status[document_id]["progress"] = progress
            self.processing_status[document_id]["message"] = stage

            if callback:
                await callback(self.processing_status[document_id])

    async def list_documents(
        self, limit: int = 10, offset: int = 0, status_filter: Optional[str] = None
    ):
        """List processed documents"""
        docs = list(self.processing_status.values())
        if status_filter:
            docs = [d for d in docs if d.get("status") == status_filter]
        return docs[offset : offset + limit]
